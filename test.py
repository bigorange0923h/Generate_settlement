from docx import Document

doc = Document('demo.docx')

# 每一段的内容
for para in doc.paragraphs:
    print(para.text)

# 每一段的编号、内容
for i in range(len(doc.paragraphs)):
    print(str(i),  doc.paragraphs[i].text)

# 表格
tbs = doc.tables
for tb in tbs:
    # 行
    for row in tb.rows:
        # 列
        for cell in row.cells:
            print(cell.text)
            # 也可以用下面方法
            '''text = ''
            for p in cell.paragraphs:
                text += p.text
            print(text)'''
